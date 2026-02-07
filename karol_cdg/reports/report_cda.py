"""
Generazione report per il Consiglio di Amministrazione (CdA).

Produce un documento Word (.docx) contenente:
  - Sintesi esecutiva con CE consolidato e KPI principali
  - Conto Economico per Unita' Operativa
  - Sezione KPI con semafori (verde/giallo/rosso)
  - Flusso di cassa (cash flow)
  - Alert attivi con livello di criticita'

Il report e' pensato per essere distribuito ai membri del CdA
con cadenza mensile.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from karol_cdg.config import (
    MESI_IT,
    UNITA_OPERATIVE,
    SOGLIE_SEMAFORO,
    SIMBOLO_VALUTA,
    LivelliAlert,
)
from karol_cdg.utils.format_utils import (
    formatta_valuta,
    formatta_percentuale,
    formatta_numero,
)
from karol_cdg.utils.date_utils import formatta_periodo_esteso

logger = logging.getLogger(__name__)

# ============================================================================
# COSTANTI STILE DOCUMENTO
# ============================================================================

_COLORE_INTESTAZIONE = RGBColor(0x1B, 0x3A, 0x6B)  # Blu scuro Karol
_COLORE_VERDE = RGBColor(0x27, 0xAE, 0x60)
_COLORE_GIALLO = RGBColor(0xF3, 0x9C, 0x12)
_COLORE_ROSSO = RGBColor(0xE7, 0x4C, 0x3C)
_COLORE_GRIGIO = RGBColor(0x7F, 0x8C, 0x8D)

_MAPPA_COLORI_SEMAFORO = {
    "verde": _COLORE_VERDE,
    "giallo": _COLORE_GIALLO,
    "rosso": _COLORE_ROSSO,
}

_NOME_AZIENDA = "Gruppo Karol S.p.A."
_TITOLO_REPORT = "Report Consiglio di Amministrazione"


# ============================================================================
# FUNZIONE PRINCIPALE
# ============================================================================


def genera_report_cda(
    periodo: str,
    ce_consolidato: dict,
    ce_per_uo: dict,
    kpi: list,
    cash_flow: dict,
    alert: list,
    output_path: Path,
) -> Path:
    """
    Genera il report completo per il Consiglio di Amministrazione in formato Word.

    Il documento viene strutturato nelle seguenti sezioni:
      1. Copertina e intestazione
      2. Sintesi esecutiva (CE consolidato + KPI principali)
      3. Conto Economico per Unita' Operativa
      4. Indicatori KPI con semafori
      5. Flusso di cassa
      6. Alert e segnalazioni

    Parametri
    ---------
    periodo : str
        Periodo di riferimento nel formato "MM/YYYY", es. "01/2026".
    ce_consolidato : dict
        Dati del Conto Economico consolidato. Attese chiavi:
        - "ricavi_totali" (float)
        - "costi_diretti_totali" (float)
        - "mol_industriale" (float)
        - "costi_sede" (float)
        - "mol_gestionale" (float)
        - "risultato_netto" (float)
        - "voci" (list[dict]) con dettaglio voci
    ce_per_uo : dict
        Dizionario {codice_uo: dati_ce} con CE per ciascuna U.O.
    kpi : list
        Lista di dizionari KPI, ciascuno con:
        - "nome" (str), "valore" (float), "formato" (str),
        - "semaforo" (str: "verde"/"giallo"/"rosso"),
        - "target" (float, opzionale), "nota" (str, opzionale)
    cash_flow : dict
        Dati del flusso di cassa. Attese chiavi:
        - "saldo_iniziale" (float)
        - "entrate" (float)
        - "uscite" (float)
        - "saldo_finale" (float)
        - "dettaglio" (list[dict], opzionale) righe dettaglio
    alert : list
        Lista di dizionari alert con:
        - "livello" (str: "verde"/"giallo"/"rosso")
        - "area" (str), "messaggio" (str)
    output_path : Path
        Percorso del file di output .docx.

    Ritorna
    -------
    Path
        Percorso del file Word generato.

    Raises
    ------
    IOError
        Se non e' possibile scrivere il file di output.
    """
    logger.info("Generazione report CdA per il periodo %s", periodo)

    # Assicura che la directory di output esista
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Imposta margini del documento
    for sezione in doc.sections:
        sezione.top_margin = Cm(2)
        sezione.bottom_margin = Cm(2)
        sezione.left_margin = Cm(2.5)
        sezione.right_margin = Cm(2.5)

    # --- Copertina ---
    _aggiungi_copertina(doc, periodo)

    # --- Sezione 1: Sintesi esecutiva ---
    doc.add_page_break()
    _sezione_sintesi_esecutiva(doc, periodo, ce_consolidato, kpi)

    # --- Sezione 2: CE per U.O. ---
    doc.add_page_break()
    _sezione_ce_per_uo(doc, ce_per_uo, periodo)

    # --- Sezione 3: KPI ---
    doc.add_page_break()
    _sezione_kpi(doc, kpi)

    # --- Sezione 4: Cash Flow ---
    doc.add_page_break()
    _sezione_cash_flow(doc, cash_flow)

    # --- Sezione 5: Alert ---
    doc.add_page_break()
    _sezione_alert(doc, alert)

    # --- Pie' di pagina ---
    _aggiungi_pie_di_pagina(doc, periodo)

    # Salva il documento
    try:
        doc.save(str(output_path))
        logger.info("Report CdA salvato in: %s", output_path)
    except Exception as exc:
        logger.error("Errore nel salvataggio del report CdA: %s", exc)
        raise IOError(f"Impossibile salvare il report CdA in {output_path}: {exc}")

    return output_path


# ============================================================================
# SEZIONI DEL REPORT
# ============================================================================


def _sezione_sintesi_esecutiva(
    doc: Document,
    periodo: str,
    ce_consolidato: dict,
    kpi: list,
) -> None:
    """
    Aggiunge la sezione di sintesi esecutiva al documento.

    Contiene una panoramica dei risultati economici consolidati e
    gli indicatori chiave piu' rilevanti per il CdA.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    periodo : str
        Periodo di riferimento "MM/YYYY".
    ce_consolidato : dict
        Dati del CE consolidato (vedi genera_report_cda).
    kpi : list
        Lista KPI principali (vedi genera_report_cda).
    """
    _formatta_intestazione(doc, "1. Sintesi Esecutiva", livello=1)

    periodo_esteso = formatta_periodo_esteso(periodo)
    paragrafo_intro = doc.add_paragraph()
    paragrafo_intro.add_run(
        f"Di seguito la sintesi dei risultati economici consolidati "
        f"del {_NOME_AZIENDA} per il mese di {periodo_esteso}."
    )

    # Tabella riepilogo CE consolidato
    _formatta_intestazione(doc, "Riepilogo Conto Economico Consolidato", livello=2)

    ricavi = ce_consolidato.get("ricavi_totali", 0.0)
    costi_diretti = ce_consolidato.get("costi_diretti_totali", 0.0)
    mol_industriale = ce_consolidato.get("mol_industriale", 0.0)
    costi_sede = ce_consolidato.get("costi_sede", 0.0)
    mol_gestionale = ce_consolidato.get("mol_gestionale", 0.0)
    risultato_netto = ce_consolidato.get("risultato_netto", 0.0)

    # Calcola percentuali su ricavi (evita divisione per zero)
    pct_costi_diretti = (costi_diretti / ricavi * 100) if ricavi else 0.0
    pct_mol_ind = (mol_industriale / ricavi * 100) if ricavi else 0.0
    pct_costi_sede = (costi_sede / ricavi * 100) if ricavi else 0.0
    pct_mol_gest = (mol_gestionale / ricavi * 100) if ricavi else 0.0

    intestazioni_ce = ["Voce", "Importo", "% su Ricavi"]
    righe_ce = [
        ["Ricavi totali", formatta_valuta(ricavi), "100,0%"],
        ["Costi diretti totali", formatta_valuta(costi_diretti),
         formatta_percentuale(pct_costi_diretti)],
        ["MOL Industriale", formatta_valuta(mol_industriale),
         formatta_percentuale(pct_mol_ind)],
        ["Costi sede allocati", formatta_valuta(costi_sede),
         formatta_percentuale(pct_costi_sede)],
        ["MOL Gestionale", formatta_valuta(mol_gestionale),
         formatta_percentuale(pct_mol_gest)],
        ["Risultato netto", formatta_valuta(risultato_netto), ""],
    ]
    _aggiungi_tabella(doc, intestazioni_ce, righe_ce)

    # KPI principali in sintesi (primi 5)
    if kpi:
        _formatta_intestazione(doc, "Indicatori Chiave", livello=2)
        kpi_principali = kpi[:5]
        intestazioni_kpi = ["Indicatore", "Valore", "Stato"]
        righe_kpi = []
        for indicatore in kpi_principali:
            nome = indicatore.get("nome", "")
            valore = indicatore.get("valore", 0.0)
            formato = indicatore.get("formato", "numero")
            semaforo = indicatore.get("semaforo", "")

            # Formatta il valore in base al tipo
            if formato == "valuta":
                valore_fmt = formatta_valuta(valore)
            elif formato == "percentuale":
                valore_fmt = formatta_percentuale(valore)
            else:
                valore_fmt = formatta_numero(valore)

            # Simbolo testuale del semaforo
            simbolo_semaforo = _simbolo_semaforo(semaforo)
            righe_kpi.append([nome, valore_fmt, simbolo_semaforo])

        _aggiungi_tabella(doc, intestazioni_kpi, righe_kpi)

    logger.debug("Sezione sintesi esecutiva aggiunta al documento.")


def _sezione_ce_per_uo(doc: Document, ce_per_uo: dict, periodo: str) -> None:
    """
    Aggiunge la sezione del Conto Economico per singola Unita' Operativa.

    Per ogni U.O. presente nei dati viene creata una sotto-sezione con
    una tabella riassuntiva delle voci di conto economico.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    ce_per_uo : dict
        Dizionario {codice_uo: dati_ce}. Ogni dati_ce contiene:
        - "ricavi" (float)
        - "costi_diretti" (float)
        - "mol_industriale" (float)
        - "costi_sede" (float)
        - "mol_gestionale" (float)
        - "voci" (list[dict]) con "codice", "descrizione", "importo"
    periodo : str
        Periodo di riferimento "MM/YYYY".
    """
    _formatta_intestazione(doc, "2. Conto Economico per Unita' Operativa", livello=1)

    periodo_esteso = formatta_periodo_esteso(periodo)
    doc.add_paragraph(
        f"Dettaglio del Conto Economico Industriale per ciascuna Unita' "
        f"Operativa nel periodo di {periodo_esteso}."
    )

    if not ce_per_uo:
        doc.add_paragraph(
            "Nessun dato disponibile per il periodo selezionato.",
        ).italic = True
        logger.warning("Nessun CE per U.O. disponibile per il periodo %s", periodo)
        return

    # Ordina le U.O. per codice
    for codice_uo in sorted(ce_per_uo.keys()):
        dati_uo = ce_per_uo[codice_uo]

        # Recupera il nome dell'U.O. dall'anagrafica
        uo_info = UNITA_OPERATIVE.get(codice_uo)
        nome_uo = uo_info.nome if uo_info else codice_uo

        _formatta_intestazione(doc, f"{codice_uo} - {nome_uo}", livello=2)

        # Tabella riepilogativa
        ricavi = dati_uo.get("ricavi", 0.0)
        costi_diretti = dati_uo.get("costi_diretti", 0.0)
        mol_industriale = dati_uo.get("mol_industriale", 0.0)
        costi_sede = dati_uo.get("costi_sede", 0.0)
        mol_gestionale = dati_uo.get("mol_gestionale", 0.0)

        pct_mol_ind = (mol_industriale / ricavi * 100) if ricavi else 0.0
        pct_mol_gest = (mol_gestionale / ricavi * 100) if ricavi else 0.0

        intestazioni = ["Aggregato", "Importo", "% Ricavi"]
        righe = [
            ["Ricavi", formatta_valuta(ricavi), "100,0%"],
            ["Costi diretti", formatta_valuta(costi_diretti), ""],
            ["MOL Industriale", formatta_valuta(mol_industriale),
             formatta_percentuale(pct_mol_ind)],
            ["Costi sede allocati", formatta_valuta(costi_sede), ""],
            ["MOL Gestionale", formatta_valuta(mol_gestionale),
             formatta_percentuale(pct_mol_gest)],
        ]
        _aggiungi_tabella(doc, intestazioni, righe)

        # Dettaglio voci se disponibili
        voci = dati_uo.get("voci", [])
        if voci:
            doc.add_paragraph("")  # Spaziatura
            intestazioni_voci = ["Codice", "Descrizione", "Importo"]
            righe_voci = []
            for voce in voci:
                righe_voci.append([
                    voce.get("codice", ""),
                    voce.get("descrizione", ""),
                    formatta_valuta(voce.get("importo", 0.0)),
                ])
            _aggiungi_tabella(doc, intestazioni_voci, righe_voci)

    logger.debug("Sezione CE per U.O. aggiunta al documento.")


def _sezione_kpi(doc: Document, kpi: list) -> None:
    """
    Aggiunge la sezione completa degli indicatori KPI con semafori.

    Per ogni KPI viene mostrato il valore corrente, il target (se
    disponibile), il livello semaforo e una eventuale nota esplicativa.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    kpi : list
        Lista completa dei KPI (vedi genera_report_cda).
    """
    _formatta_intestazione(doc, "3. Indicatori di Performance (KPI)", livello=1)

    doc.add_paragraph(
        "Panoramica completa degli indicatori di performance con "
        "rappresentazione a semaforo (verde = in linea, giallo = attenzione, "
        "rosso = critico)."
    )

    if not kpi:
        doc.add_paragraph(
            "Nessun indicatore KPI disponibile.",
        ).italic = True
        logger.warning("Nessun KPI disponibile per il report.")
        return

    intestazioni = ["Indicatore", "Valore", "Target", "Stato", "Note"]
    righe = []

    for indicatore in kpi:
        nome = indicatore.get("nome", "")
        valore = indicatore.get("valore", 0.0)
        formato = indicatore.get("formato", "numero")
        semaforo = indicatore.get("semaforo", "")
        target = indicatore.get("target", None)
        nota = indicatore.get("nota", "")

        # Formattazione del valore
        if formato == "valuta":
            valore_fmt = formatta_valuta(valore)
            target_fmt = formatta_valuta(target) if target is not None else "-"
        elif formato == "percentuale":
            valore_fmt = formatta_percentuale(valore)
            target_fmt = formatta_percentuale(target) if target is not None else "-"
        elif formato == "giorni":
            valore_fmt = f"{formatta_numero(valore)} gg"
            target_fmt = f"{formatta_numero(target)} gg" if target is not None else "-"
        else:
            valore_fmt = formatta_numero(valore)
            target_fmt = formatta_numero(target) if target is not None else "-"

        simbolo = _simbolo_semaforo(semaforo)

        righe.append([nome, valore_fmt, target_fmt, simbolo, nota])

    _aggiungi_tabella(doc, intestazioni, righe)

    # Riepilogo semafori
    conteggio = {"verde": 0, "giallo": 0, "rosso": 0}
    for indicatore in kpi:
        sem = indicatore.get("semaforo", "")
        if sem in conteggio:
            conteggio[sem] += 1

    totale_kpi = len(kpi)
    paragrafo_riepilogo = doc.add_paragraph()
    paragrafo_riepilogo.add_run(
        f"\nRiepilogo: {conteggio['verde']} verdi, "
        f"{conteggio['giallo']} gialli, "
        f"{conteggio['rosso']} rossi su {totale_kpi} indicatori totali."
    )

    logger.debug("Sezione KPI aggiunta al documento.")


def _sezione_cash_flow(doc: Document, cash_flow: dict) -> None:
    """
    Aggiunge la sezione del flusso di cassa al documento.

    Mostra il saldo iniziale, le entrate e uscite del periodo e il
    saldo finale, con eventuale dettaglio per voce.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    cash_flow : dict
        Dati del flusso di cassa (vedi genera_report_cda).
    """
    _formatta_intestazione(doc, "4. Flusso di Cassa", livello=1)

    if not cash_flow:
        doc.add_paragraph(
            "Nessun dato di flusso di cassa disponibile.",
        ).italic = True
        logger.warning("Nessun dato cash flow disponibile per il report.")
        return

    doc.add_paragraph(
        "Situazione del flusso di cassa per il periodo di riferimento."
    )

    # Tabella riepilogo cash flow
    saldo_iniziale = cash_flow.get("saldo_iniziale", 0.0)
    entrate = cash_flow.get("entrate", 0.0)
    uscite = cash_flow.get("uscite", 0.0)
    saldo_finale = cash_flow.get("saldo_finale", 0.0)

    intestazioni = ["Voce", "Importo"]
    righe = [
        ["Saldo iniziale", formatta_valuta(saldo_iniziale)],
        ["(+) Entrate del periodo", formatta_valuta(entrate)],
        ["(-) Uscite del periodo", formatta_valuta(uscite)],
        ["Saldo finale", formatta_valuta(saldo_finale)],
    ]
    _aggiungi_tabella(doc, intestazioni, righe)

    # Variazione netta
    variazione = saldo_finale - saldo_iniziale
    segno = "+" if variazione >= 0 else ""
    paragrafo_var = doc.add_paragraph()
    paragrafo_var.add_run(
        f"Variazione netta del periodo: {segno}{formatta_valuta(variazione)}"
    ).bold = True

    # Dettaglio per voce se disponibile
    dettaglio = cash_flow.get("dettaglio", [])
    if dettaglio:
        _formatta_intestazione(doc, "Dettaglio movimenti", livello=2)
        intestazioni_det = ["Descrizione", "Tipo", "Importo"]
        righe_det = []
        for riga in dettaglio:
            righe_det.append([
                riga.get("descrizione", ""),
                riga.get("tipo", ""),
                formatta_valuta(riga.get("importo", 0.0)),
            ])
        _aggiungi_tabella(doc, intestazioni_det, righe_det)

    logger.debug("Sezione cash flow aggiunta al documento.")


def _sezione_alert(doc: Document, alert: list) -> None:
    """
    Aggiunge la sezione degli alert e segnalazioni al documento.

    Gli alert sono raggruppati per livello di criticita' (rosso, giallo,
    verde) e mostrati con il relativo indicatore visivo.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    alert : list
        Lista degli alert attivi (vedi genera_report_cda).
    """
    _formatta_intestazione(doc, "5. Alert e Segnalazioni", livello=1)

    if not alert:
        doc.add_paragraph(
            "Nessun alert attivo nel periodo di riferimento. "
            "Tutti gli indicatori rientrano nelle soglie previste."
        )
        logger.info("Nessun alert da inserire nel report CdA.")
        return

    doc.add_paragraph(
        f"Sono presenti {len(alert)} segnalazioni attive. "
        "Di seguito il dettaglio per livello di criticita'."
    )

    # Ordina alert per livello: rosso prima, poi giallo, poi verde
    ordine_livello = {"rosso": 0, "giallo": 1, "verde": 2}
    alert_ordinati = sorted(
        alert,
        key=lambda a: ordine_livello.get(a.get("livello", "verde"), 99),
    )

    # Tabella alert
    intestazioni = ["Livello", "Area", "Descrizione"]
    righe = []
    for allarme in alert_ordinati:
        livello = allarme.get("livello", "verde")
        area = allarme.get("area", "")
        messaggio = allarme.get("messaggio", "")
        simbolo = _simbolo_semaforo(livello)
        righe.append([simbolo, area, messaggio])

    _aggiungi_tabella(doc, intestazioni, righe)

    # Conteggio per livello
    conteggio = {"rosso": 0, "giallo": 0, "verde": 0}
    for allarme in alert:
        liv = allarme.get("livello", "verde")
        if liv in conteggio:
            conteggio[liv] += 1

    paragrafo_sintesi = doc.add_paragraph()
    paragrafo_sintesi.add_run(
        f"Riepilogo alert: {conteggio['rosso']} critici, "
        f"{conteggio['giallo']} di attenzione, "
        f"{conteggio['verde']} informativi."
    )

    logger.debug("Sezione alert aggiunta al documento.")


# ============================================================================
# FUNZIONI DI SUPPORTO
# ============================================================================


def _aggiungi_tabella(
    doc: Document,
    intestazioni: list,
    righe: list,
    stile: str = "Table Grid",
) -> None:
    """
    Aggiunge una tabella formattata al documento Word.

    La prima riga contiene le intestazioni in grassetto con sfondo
    colorato. Le righe successive contengono i dati.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    intestazioni : list
        Lista di stringhe per le intestazioni delle colonne.
    righe : list
        Lista di liste di stringhe con i dati delle righe.
    stile : str
        Nome dello stile tabella Word (default: "Table Grid").
    """
    n_colonne = len(intestazioni)
    n_righe = len(righe) + 1  # +1 per la riga di intestazione

    tabella = doc.add_table(rows=n_righe, cols=n_colonne)

    # Tenta di applicare lo stile; se non disponibile usa il default
    try:
        tabella.style = stile
    except KeyError:
        logger.debug(
            "Stile tabella '%s' non disponibile, uso stile predefinito.", stile
        )

    tabella.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Riga intestazione
    riga_intestazione = tabella.rows[0]
    for idx, testo in enumerate(intestazioni):
        cella = riga_intestazione.cells[idx]
        cella.text = ""
        paragrafo = cella.paragraphs[0]
        run = paragrafo.add_run(testo)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sfondo intestazione (blu scuro)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1B3A6B")
        shading.set(qn("w:val"), "clear")
        cella._element.get_or_add_tcPr().append(shading)

    # Righe dati
    for idx_riga, dati_riga in enumerate(righe):
        riga_tabella = tabella.rows[idx_riga + 1]
        for idx_col, valore in enumerate(dati_riga):
            if idx_col < n_colonne:
                cella = riga_tabella.cells[idx_col]
                cella.text = ""
                paragrafo = cella.paragraphs[0]
                run = paragrafo.add_run(str(valore))
                run.font.size = Pt(9)

                # Allineamento: numeri e importi a destra, testo a sinistra
                if idx_col > 0 and any(
                    c in str(valore) for c in [SIMBOLO_VALUTA, "%", "gg"]
                ):
                    paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif idx_col == 0:
                    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sfondo alternato per leggibilita'
        if idx_riga % 2 == 1:
            for idx_col in range(n_colonne):
                cella = riga_tabella.cells[idx_col]
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F3F4")
                shading.set(qn("w:val"), "clear")
                cella._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph("")  # Spaziatura dopo la tabella


def _formatta_intestazione(doc: Document, testo: str, livello: int) -> None:
    """
    Aggiunge un'intestazione (heading) formattata al documento.

    Utilizza lo stile intestazione nativo di Word con personalizzazione
    del colore e della dimensione del font.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    testo : str
        Testo dell'intestazione.
    livello : int
        Livello dell'intestazione (1 = principale, 2 = sotto-sezione, ecc.).
    """
    intestazione = doc.add_heading(testo, level=livello)

    # Personalizza il colore dell'intestazione
    for run in intestazione.runs:
        run.font.color.rgb = _COLORE_INTESTAZIONE
        if livello == 1:
            run.font.size = Pt(16)
        elif livello == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)


def _aggiungi_copertina(doc: Document, periodo: str) -> None:
    """
    Aggiunge la pagina di copertina al documento.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    periodo : str
        Periodo di riferimento "MM/YYYY".
    """
    periodo_esteso = formatta_periodo_esteso(periodo)
    data_generazione = datetime.now().strftime("%d/%m/%Y %H:%M")

    # Spaziatura superiore
    for _ in range(6):
        doc.add_paragraph("")

    # Titolo azienda
    paragrafo_azienda = doc.add_paragraph()
    paragrafo_azienda.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_azienda = paragrafo_azienda.add_run(_NOME_AZIENDA)
    run_azienda.bold = True
    run_azienda.font.size = Pt(24)
    run_azienda.font.color.rgb = _COLORE_INTESTAZIONE

    doc.add_paragraph("")  # Spaziatura

    # Titolo report
    paragrafo_titolo = doc.add_paragraph()
    paragrafo_titolo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titolo = paragrafo_titolo.add_run(_TITOLO_REPORT)
    run_titolo.bold = True
    run_titolo.font.size = Pt(20)
    run_titolo.font.color.rgb = _COLORE_GRIGIO

    doc.add_paragraph("")  # Spaziatura

    # Periodo
    paragrafo_periodo = doc.add_paragraph()
    paragrafo_periodo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_periodo = paragrafo_periodo.add_run(periodo_esteso)
    run_periodo.font.size = Pt(18)
    run_periodo.font.color.rgb = _COLORE_INTESTAZIONE

    # Spaziatura inferiore
    for _ in range(8):
        doc.add_paragraph("")

    # Data di generazione
    paragrafo_data = doc.add_paragraph()
    paragrafo_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_data = paragrafo_data.add_run(
        f"Documento generato il {data_generazione}"
    )
    run_data.font.size = Pt(9)
    run_data.font.color.rgb = _COLORE_GRIGIO
    run_data.italic = True

    # Classificazione
    paragrafo_class = doc.add_paragraph()
    paragrafo_class.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_class = paragrafo_class.add_run("RISERVATO - USO INTERNO CdA")
    run_class.bold = True
    run_class.font.size = Pt(10)
    run_class.font.color.rgb = _COLORE_ROSSO


def _aggiungi_pie_di_pagina(doc: Document, periodo: str) -> None:
    """
    Aggiunge un paragrafo finale con informazioni di chiusura.

    Parametri
    ---------
    doc : Document
        Documento Word in fase di costruzione.
    periodo : str
        Periodo di riferimento "MM/YYYY".
    """
    doc.add_page_break()
    paragrafo = doc.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(
        f"--- Fine del Report CdA - {formatta_periodo_esteso(periodo)} ---\n"
        f"{_NOME_AZIENDA} - Sistema di Controllo di Gestione Karol CDG"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = _COLORE_GRIGIO
    run.italic = True


def _simbolo_semaforo(livello: str) -> str:
    """
    Restituisce un simbolo testuale per il semaforo.

    Utilizzato nelle tabelle Word dove i colori non sono facilmente
    applicabili nelle celle di testo.

    Parametri
    ---------
    livello : str
        Livello del semaforo: "verde", "giallo" o "rosso".

    Ritorna
    -------
    str
        Simbolo testuale: "[VERDE]", "[GIALLO]" o "[ROSSO]".
    """
    mappa = {
        "verde": "[VERDE]",
        "giallo": "[GIALLO]",
        "rosso": "[ROSSO]",
    }
    return mappa.get(livello, "[N/D]")
